from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "小智分层伴学知识库.md"
OUTPUT = ROOT / "小智分层伴学知识库.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\86153\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.623.12021\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


FONT = "Microsoft YaHei"
GREEN = RGBColor(46, 125, 50)
BLUE = RGBColor(46, 116, 181)
ORANGE = RGBColor(197, 106, 26)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(95, 99, 104)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_style(style, size, color, before, after, line=1.25, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, 9, color=GRAY)


def stage_color(text):
    if "[小学]" in text:
        return GREEN
    if "[初中]" in text:
        return BLUE
    if "[高中]" in text:
        return ORANGE
    return BLUE


def clean_inline(text):
    return text.replace("`", "").replace("  ", " ").strip()


def add_body(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(clean_inline(text))
    set_font(r, 10.5)
    return p


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(clean_inline(value))
            set_font(r, 8.5, bold=(i == 0), color=DARK if i == 0 else None)
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E8EEF5")
                cell._tc.get_or_add_tcPr().append(shd)
    if cols == 5:
        widths = [1200, 2040, 2040, 2040, 2040]
    elif cols == 2:
        widths = [2700, 6660]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    configure_style(styles["Normal"], 10.5, RGBColor(30, 30, 30), 0, 6)
    configure_style(styles["Title"], 26, GREEN, 0, 8, 1.0, True)
    configure_style(styles["Heading 1"], 16, BLUE, 18, 10, 1.1, True)
    configure_style(styles["Heading 2"], 13, BLUE, 14, 7, 1.1, True)
    configure_style(styles["Heading 3"], 12, DARK, 10, 5, 1.1, True)
    for name in ("List Bullet", "List Number"):
        configure_style(styles[name], 10.5, RGBColor(30, 30, 30), 0, 4)
        styles[name].paragraph_format.left_indent = Inches(0.375)
        styles[name].paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("小智分层伴学知识库"), 9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    first_heading = True
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip() == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-|]+\|$", lines[i + 1]):
            rows = []
            i += 2
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            text = clean_inline(line[2:])
            if first_heading:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(72)
                p.paragraph_format.space_after = Pt(12)
                set_font(p.add_run(text), 26, True, GREEN)
                first_heading = False
            else:
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.keep_with_next = True
                set_font(p.add_run(text), 16, True, stage_color(text))
        elif line.startswith("## "):
            text = clean_inline(line[3:])
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.keep_with_next = True
            set_font(p.add_run(text), 13, True, stage_color(text))
        elif re.match(r"^- ", line):
            p = add_body(doc, line[2:], "List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
        elif re.match(r"^\d+\. ", line):
            p = add_body(doc, re.sub(r"^\d+\. ", "", line), "List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
        elif line.strip():
            p = add_body(doc, line)
            if line.startswith(("问：", "答：", "检索词：", "唯一答案：")):
                p.paragraph_format.space_after = Pt(4)
                if line.startswith(("检索词：", "唯一答案：")):
                    for run in p.runs:
                        run.bold = True
        i += 1

    props = doc.core_properties
    props.title = "小智分层伴学知识库"
    props.subject = "分层学习问答、方法指导和专注提醒"
    props.author = "小智伴学项目"
    props.keywords = "小学,初中,高中,语文,数学,英语,学习方法,专注提醒"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
